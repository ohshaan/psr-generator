import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
from datetime import datetime
import chardet
import zipfile
import sys
import os

# ---------- Helper Functions ----------

def robust_read_file(uploaded_file):
    """Read CSV or Excel with encoding fallback."""
    name = uploaded_file.name.lower()
    if name.endswith((".xls", ".xlsx")):
        return pd.read_excel(uploaded_file)
    for enc in ('utf-8', 'latin1'):
        try:
            uploaded_file.seek(0)
            return pd.read_csv(uploaded_file, encoding=enc)
        except Exception:
            continue
    uploaded_file.seek(0)
    raw = uploaded_file.read()
    detected = chardet.detect(raw)['encoding'] or 'utf-8'
    uploaded_file.seek(0)
    return pd.read_csv(uploaded_file, encoding=detected)

def group_activities(df_proj):
    """Group & dedupe activities by Task Reference Number in date order."""
    df = df_proj.copy()
    df['Date'] = pd.to_datetime(df['Date'], errors='coerce')
    # Avoid using Series.replace with an empty string because it treats the
    # pattern as a regular expression which would replace the empty string in
    # every position of every value. Instead strip/clean first and then update
    # only the rows that are actually empty.
    df['Task Reference Number'] = (
        df.get('Task Reference Number', '')
          .fillna('')
          .astype(str)
          .str.strip()
    )
    df.loc[df['Task Reference Number'] == '', 'Task Reference Number'] = (
        'No Reference'
    )
    buckets = {}
    for _, row in df.iterrows():
        task_ref = row['Task Reference Number']
        details = (row.get('Modification Details', '') or row.get('Task Description', '')).strip()
        date = row['Date']
        date_str = date.strftime('%d-%b-%Y') if pd.notnull(date) else ''
        extras = []
        for col in ('Remarks', 'Notes'):
            val = str(row.get(col, '') or '').strip()
            if val and val.lower() != 'nan':
                extras.append(f"{col}: {val}")
        parts = [p for p in (date_str, details) if p]
        if extras:
            parts += extras
        activity = " | ".join(parts)
        if not activity:
            continue
        key = (date if pd.notnull(date) else pd.Timestamp.max, activity)
        buckets.setdefault(task_ref, set()).add(key)

    grouped = {}
    for task_ref, items in buckets.items():
        sorted_items = sorted(items, key=lambda x: x[0])
        seen, lines = set(), []
        for _, txt in sorted_items:
            if txt not in seen:
                seen.add(txt)
                lines.append(txt)
        grouped[task_ref] = lines

    return grouped

def replace_placeholders_preserve_format(doc, mapping):
    """
    For every paragraph in the document (body, tables, footers):
      1) Build a flat list of (char, format_dict) from all runs.
      2) Search the combined text for any "{Key}" in mapping.
      3) Whenever a placeholder is found, replace it with mapping[Key], and
         assign the replacement text the same run-format as the first character
         of the placeholder. Everything else keeps its original format.
      4) Rebuild the paragraph’s runs by grouping consecutive chars that share
         the same format_dict.
    """

    def get_run_format(run):
        font = run.font
        return {
            'name': font.name,
            'size': font.size,
            'color': font.color.rgb,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline if hasattr(run, 'underline') else None,
        }

    def apply_format_to_run(run, fmt):
        if fmt['name'] is not None:
            run.font.name = fmt['name']
        if fmt['size'] is not None:
            run.font.size = fmt['size']
        if fmt['color'] is not None:
            run.font.color.rgb = fmt['color']
        if fmt.get('bold') is not None:
            run.bold = fmt['bold']
        if fmt.get('italic') is not None:
            run.italic = fmt['italic']
        if fmt.get('underline') is not None:
            run.underline = fmt['underline']

    def replace_in_para(para):
        # 1) Build a list of (char, fmt) for every character in all runs
        char_fmt = []
        for run in para.runs:
            fmt = get_run_format(run)
            for ch in run.text:
                char_fmt.append((ch, fmt))

        full_text = ''.join(ch for ch, _ in char_fmt)
        if not full_text:
            return  # nothing in this paragraph

        # 2) Locate all placeholder occurrences (start, end, replacement_text)
        replacements = []
        for key, val in mapping.items():
            placeholder = f"{{{key}}}"
            start = 0
            while True:
                idx = full_text.find(placeholder, start)
                if idx == -1:
                    break
                replacements.append((idx, idx + len(placeholder), str(val)))
                start = idx + len(placeholder)

        if not replacements:
            return  # no "{Key}" in this paragraph

        # 3) Sort by start index
        replacements.sort(key=lambda x: x[0])

        # 4) Build a new char_fmt list with replacements in place
        new_char_fmt = []
        cursor = 0
        for start, end, repl_text in replacements:
            # a) Copy everything before this placeholder
            new_char_fmt.extend(char_fmt[cursor:start])

            # b) Copy the replacement text with the same format as first placeholder char
            if start < len(char_fmt):
                repl_fmt = char_fmt[start][1]
            else:
                repl_fmt = {
                    'name': None, 'size': None, 'color': None,
                    'bold': None, 'italic': None, 'underline': None
                }
            for ch in repl_text:
                new_char_fmt.append((ch, repl_fmt))

            cursor = end

        # c) Copy anything after the last placeholder
        new_char_fmt.extend(char_fmt[cursor:])

        # 5) Remove all existing runs
        for _ in range(len(para.runs)):
            para.runs[0]._element.getparent().remove(para.runs[0]._element)

        # 6) Rebuild runs by grouping consecutive chars that share the same fmt
        if not new_char_fmt:
            return

        group_chars = [new_char_fmt[0][0]]
        group_fmt = new_char_fmt[0][1]
        for ch, fmt in new_char_fmt[1:]:
            if fmt == group_fmt:
                group_chars.append(ch)
            else:
                run = para.add_run(''.join(group_chars))
                apply_format_to_run(run, group_fmt)
                group_chars = [ch]
                group_fmt = fmt

        run = para.add_run(''.join(group_chars))
        apply_format_to_run(run, group_fmt)

    # ==== Apply to every paragraph in the main body ====
    for para in doc.paragraphs:
        replace_in_para(para)

    # ==== Apply inside every table cell ====
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)

    # ==== Apply inside each footer paragraph ====
    for section in doc.sections:
        for para in section.footer.paragraphs:
            replace_in_para(para)

    return doc

def set_footer_for_all_sections(doc, footer_text):
    for section in doc.sections:
        footer = section.footer
        # Remove all existing paragraphs in the footer
        p_elements = list(footer._element.xpath('.//w:p'))
        for p in p_elements:
            p.getparent().remove(p)
        # Add new footer paragraph
        para = footer.add_paragraph(footer_text)
        # (Optional) Style the footer here, e.g., para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return doc

def insert_grouped_activities(doc, grouped):
    """
    Inject grouped activities wherever the placeholder {Activities} appears,
    including inside table cells. For each task reference, it writes a bold header
    followed by bullet‐pointed activities.
    """
    bullet = '•'
    font_size = Pt(10)
    font_name = 'Calibri'
    font_color = RGBColor(0, 0, 0)

    def style_run(run):
        run.font.size = font_size
        run.font.name = font_name
        run.font.color.rgb = font_color

    def style_paragraph(paragraph, make_bold=False):
        for run in paragraph.runs:
            style_run(run)
            if make_bold:
                run.bold = True

    replaced = False

    # 1) Scan all table cells for "{Activities}"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if any("{Activities}" in run.text for para in cell.paragraphs for run in para.runs):
                    # Clear out everything in that cell
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.clear()

                    # Insert header + bullets for each task_ref
                    for task_ref, activities in grouped.items():
                        if task_ref != 'No Reference':
                            p_header = cell.add_paragraph(task_ref)
                            style_paragraph(p_header, make_bold=True)

                        for act_line in activities:
                            p_act = cell.add_paragraph(f"{bullet} {act_line}")
                            style_paragraph(p_act)

                    replaced = True

    # 2) If no cell placeholder was found, append at end
    if not replaced and grouped:
        hdr = doc.add_paragraph("Recent Activities")
        hdr.runs[0].bold = True
        for task_ref, activities in grouped.items():
            if task_ref != 'No Reference':
                p = doc.add_paragraph(task_ref)
                p.runs[0].bold = True
                style_paragraph(p, make_bold=True)
            for act_line in activities:
                p_act = doc.add_paragraph(f"{bullet} {act_line}")
                style_paragraph(p_act)

    return doc

def get_footer(project_name):
    today = datetime.now().strftime('%Y-%m-%d')
    return f"Project Status Report  –   / {project_name} - {today}"

def generate_psr_report(df_proj, tpl_bytes):
    first = df_proj.iloc[0]
    proj_name = first['Project Name']
    df_proj['Date'] = pd.to_datetime(df_proj['Date'], errors='coerce')
    start, end = df_proj['Date'].min(), df_proj['Date'].max()
    days = (end - start).days + 1 if pd.notnull(start) else 0
    mapping = {
        'Project Code': first.get('Project Code', ''),
        'Project Name': proj_name,
        'Overall Status': first.get('Overall Status', ''),
        'Report Date': datetime.now().strftime('%d-%b-%Y'),
        'Days Elapsed': days,
        'Total Hours': df_proj['Hours'].sum(),
        'Report Period': f"{start:%d-%b-%Y} to {end:%d-%b-%Y}" if pd.notnull(start) else '',
        'Highlights': first.get('Highlights', ''),
        'Milestones': first.get('Milestones', ''),
        'Footer': get_footer(proj_name),
    }
    doc = Document(BytesIO(tpl_bytes))
    doc = replace_placeholders_preserve_format(doc, mapping)
    # (Do NOT call set_footer_for_all_sections here!)
    grouped = group_activities(df_proj)
    doc = insert_grouped_activities(doc, grouped)
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# ---------- Streamlit UI ----------

st.set_page_config(page_title='PSR Generator', layout='wide')
st.title('Project Status Report Generator')

# Load Word template
base_path = getattr(sys, '_MEIPASS', os.path.abspath('.'))
template_path = os.path.join(base_path, 'template.docx')
with open(template_path, 'rb') as f:
    tpl_bytes = f.read()

# File upload
uploaded = st.file_uploader('Upload Excel/CSV', type=['csv', 'xlsx'])
if not uploaded:
    st.info('Please upload your data file.')
    st.stop()

# Read data
try:
    df = robust_read_file(uploaded)
except Exception as e:
    st.error(f'Failed to read file: {e}')
    st.stop()

if df.empty:
    st.error('Uploaded file is empty.')
    st.stop()

# Column validation
required = {'Employee Name', 'Hours', 'Task Description', 'Modification Details', 'Project Name', 'Date'}
missing = required - set(df.columns)
if missing:
    st.error(f"Missing required columns: {', '.join(missing)}")
    st.stop()

# Project selection
projects = df['Project Name'].dropna().unique().tolist()
sel = st.multiselect('Select Project(s)', projects, default=projects[:1])

# Show details
for proj in sel:
    sub = df[df['Project Name'] == proj]
    st.subheader(proj)
    # Summary table
    summary = (
        sub.groupby('Employee Name')['Hours']
           .sum()
           .reset_index()
           .rename(columns={'Hours': 'Total Hours'})
           .sort_values('Total Hours', ascending=False)
    )
    st.dataframe(summary, use_container_width=True, height=200)

    # Tabs for each employee
    emp_names = summary['Employee Name'].tolist()
    tabs = st.tabs(emp_names)
    for tab, emp in zip(tabs, emp_names):
        with tab:
            emp_df = sub[sub['Employee Name'] == emp].copy()
            modified_emp_df = emp_df.copy()
            modified_emp_df['Date'] = pd.to_datetime(modified_emp_df['Date'], errors='coerce').dt.strftime('%d-%b-%Y')
            cols = [c for c in ['Date', 'Modification Details', 'Task Description', 'Remarks', 'Notes', 'Hours'] if c in modified_emp_df.columns]
            st.dataframe(modified_emp_df[cols].sort_values('Date'), use_container_width=True)

# Download buttons
if sel:
    if len(sel) == 1:
        proj = sel[0]
        dfp = df[df['Project Name'] == proj]
        report = generate_psr_report(dfp, tpl_bytes)
        safe = ''.join(c for c in proj if c.isalnum() or c in (' ', '_', '-')).rstrip()
        st.download_button(
            label='Download Report',
            data=report.getvalue(),
            file_name=f"Project Status Report_{safe}_{datetime.now():%Y%m%d}.docx",
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    else:
        buf = BytesIO()
        report_cache = {}
        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            for proj in sel:
                if proj not in report_cache:
                    dfp = df[df['Project Name'] == proj]
                    report_cache[proj] = generate_psr_report(dfp, tpl_bytes)
                report = report_cache[proj]
                safe = ''.join(c for c in proj if c.isalnum() or c in (' ', '_', '-')).rstrip()
                zf.writestr(f"{safe}_{datetime.now():%Y%m%d}.docx", report.getvalue())
        buf.seek(0)
        st.download_button(
            label='Download ZIP of Reports',
            data=buf.getvalue(),
            file_name=f"PSR_{datetime.now():%Y%m%d_%H%M%S}.zip",
            mime='application/zip'
        )