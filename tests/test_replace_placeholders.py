import pytest
from docx import Document
from app import replace_placeholders_preserve_format


def test_replace_multiple_placeholders_in_paragraph():
    doc = Document()
    para = doc.add_paragraph("Hello {Name}, again {Name}!")
    replace_placeholders_preserve_format(doc, {"Name": "World"})
    assert doc.paragraphs[0].text == "Hello World, again World!"
